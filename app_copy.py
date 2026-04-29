import os
import json
import hashlib
import time
from typing import List
from datetime import datetime, timedelta

import streamlit as st
import pandas as pd
import requests

from ortools.constraint_solver import routing_enums_pb2
from ortools.constraint_solver import pywrapcp

try:
    from github import Github
    GITHUB_AVAILABLE = True
except ImportError:
    GITHUB_AVAILABLE = False


# =========================
# Konfiguracja aplikacji
# =========================
st.set_page_config(page_title="Planowanie tras Plantpol", layout="wide")

# =========================
# Logowanie PIN
# =========================
CORRECT_PIN = st.secrets["APP_PIN"]

if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = False

def show_login():
    st.markdown("## Logowanie")
    st.write("Podaj 6-cyfrowy PIN, aby uruchomić aplikację.")

    with st.form("login_form"):
        pin_input = st.text_input("PIN", type="password", max_chars=6)
        submitted = st.form_submit_button("Zaloguj")

    if submitted:
        if pin_input == CORRECT_PIN and pin_input.isdigit() and len(pin_input) == 6:
            st.session_state["authenticated"] = True
            st.rerun()
        else:
            st.error("Nieprawidłowy PIN. Wprowadź poprawny 6-cyfrowy kod.")

if not st.session_state["authenticated"]:
    show_login()
    st.stop()


col1, col2 = st.columns([1, 5])

with col1:
    st.image("Plantpol_logo.jpg", width=180)

with col2:
    st.markdown(
        "<h1 style='margin-top:20px;'>Planowanie tras Plantpol</h1>",
        unsafe_allow_html=True
    )
st.markdown("---")

API_KEY = st.secrets["GOOGLE_MAPS_API_KEY"]

# Limit testowy (punkty)
MAX_POINTS = 150

# Cache par macierzy (punkt→punkt)
DM_PAIRS_CACHE_PATH = "dm_pairs_cache.json"

# Cache dla geocodingu
GEOCODING_CACHE_DIR = "cache_geocoding"
os.makedirs(GEOCODING_CACHE_DIR, exist_ok=True)

# CSV cache dla Streamlit Cloud
GEOCODING_CSV_PATH = "geocoding_cache.csv"


# =========================
# Helper: Geocoding GitHub cache
# =========================
def load_pairs_cache_from_github() -> dict:
    """Załaduj cache par dystansów z GitHub"""
    if not GITHUB_AVAILABLE:
        return {}
    if "GITHUB_TOKEN" not in st.secrets or "GITHUB_REPO" not in st.secrets:
        return {}
    try:
        g = Github(st.secrets["GITHUB_TOKEN"])
        repo = g.get_repo(st.secrets["GITHUB_REPO"])
        file = repo.get_contents(DM_PAIRS_CACHE_PATH)
        return json.loads(file.decoded_content.decode("utf-8"))
    except Exception:
        return {}


def load_geocoding_from_github():
    """Załaduj cache geocodingu z GitHub (minimalizuj zapytania do API)"""
    if not GITHUB_AVAILABLE:
        return None
    
    if "GITHUB_TOKEN" not in st.secrets or "GITHUB_REPO" not in st.secrets:
        return None
    
    try:
        g = Github(st.secrets["GITHUB_TOKEN"])
        repo = g.get_repo(st.secrets["GITHUB_REPO"])
        
        try:
            file = repo.get_contents(GEOCODING_CSV_PATH)
            csv_content = file.decoded_content.decode("utf-8")
            from io import StringIO
            return pd.read_csv(StringIO(csv_content))
        except:
            return None
    except:
        return None


# Inicjalizuj session_state do śledzenia nowych adresów
if "geocoding_updates" not in st.session_state:
    st.session_state["geocoding_updates"] = set()

if "geocoding_cache_df" not in st.session_state:
    # 1. Załaduj z GitHub (jeśli dostępny)
    github_cache = load_geocoding_from_github()
    if github_cache is not None and not github_cache.empty:
        st.session_state["geocoding_cache_df"] = github_cache
    # 2. Załaduj lokalny CSV (jeśli GitHub niedostępny)
    elif os.path.exists(GEOCODING_CSV_PATH):
        st.session_state["geocoding_cache_df"] = pd.read_csv(GEOCODING_CSV_PATH)
    # 3. Utwórz pusty DataFrame
    else:
        st.session_state["geocoding_cache_df"] = pd.DataFrame(
            columns=["address", "lat", "lng", "formatted_address", "status", "cached_at"]
        )

if "dm_pairs_cache" not in st.session_state:
    if os.path.exists(DM_PAIRS_CACHE_PATH):
        with open(DM_PAIRS_CACHE_PATH, "r", encoding="utf-8") as f:
            st.session_state["dm_pairs_cache"] = json.load(f)
    else:
        gh_pairs = load_pairs_cache_from_github()
        st.session_state["dm_pairs_cache"] = gh_pairs
        if gh_pairs:
            with open(DM_PAIRS_CACHE_PATH, "w", encoding="utf-8") as f:
                json.dump(gh_pairs, f)

# Baza (depot)
BASE_NAME = "_Plantpol baza"
BASE_ADDRESS = "32-600 Zaborze"
BASE_LAT = 50.0216163
BASE_LNG = 19.2408601

# Parametry biznesowe
DAY_START_HOUR = 8               # do wyświetlania godzin


# =========================
# Upload plików
# =========================

col_u1, col_u2 = st.columns(2)
with col_u1:
    points_file = st.file_uploader(
        "1) Wgraj plik z punktami (CSV ; lub XLSX) – kolumny: Kierunek;Kod;Miejscowosc;...;Ilość wózków",
        type=["csv", "xlsx"],
        key="points",
    )

with col_u2:
    vehicles_file = st.file_uploader(
        "2) Wgraj plik z pojazdami (CSV/TSV/XLSX) – kolumny: samochód, ilość wózków",
        type=["csv", "xlsx", "txt"],
        key="vehicles",
    )

from io import BytesIO
from docx import Document

# =========================
# Szablony plików do pobrania
# =========================
st.markdown("### Szablony plików")

tpl_col1, tpl_col2 = st.columns(2)

# szablon punktów
points_template_df = pd.DataFrame([
    {
        "Kod": "00-001",
        "Miejscowosc": "Warszawa",
        "Skrot_kontrahenta": "Klient A",
        "Ilość wózków": 3,
    },
    {
        "Kod": "30-002",
        "Miejscowosc": "Kraków",
        "Skrot_kontrahenta": "Klient B",
        "Ilość wózków": 2,
    },
])

# szablon pojazdów
vehicles_template_df = pd.DataFrame([
    {
        "samochód": "1",
        "ilość wózków": 11,
    },
    {
        "samochód": "2",
        "ilość wózków": 8,
    },
])

def df_to_xlsx_bytes(df: pd.DataFrame, sheet_name: str = "Arkusz1") -> bytes:
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)
    return output.getvalue()

points_template_xlsx = df_to_xlsx_bytes(points_template_df, sheet_name="Punkty")
vehicles_template_xlsx = df_to_xlsx_bytes(vehicles_template_df, sheet_name="Pojazdy")

with tpl_col1:
    st.download_button(
        "Pobierz szablon pliku punktów (.xlsx)",
        data=points_template_xlsx,
        file_name="szablon_punkty.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
    )

with tpl_col2:
    st.download_button(
        "Pobierz szablon pliku pojazdów (.xlsx)",
        data=vehicles_template_xlsx,
        file_name="szablon_pojazdy.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
    )

st.markdown("### Parametry planowania")

col_p1, col_p2, col_p3, col_p4, col_p5, col_p6 = st.columns(6)

with col_p1:
    vehicle_fixed_cost_ui = st.number_input(
        "Cena uruchomienia pojazdu",
        min_value=0,
        value=10000,
        step=1000,
    )

with col_p2:
    service_time_h = st.number_input(
        "Czas rozładunku (godziny)",
        min_value=0.0,
        value=1.5,
        step=0.5,
    )

with col_p3:
    max_route_hours = st.number_input(
        "Czas otwarcia punktów (godziny)",
        min_value=1.0,
        value=8.0,
        step=0.5,
    )

with col_p4:
    max_stops_per_route_ui = st.number_input(
        "Maks. liczba punktów na trasie",
        min_value=1,
        value=5,
        step=1,
    )

with col_p5:
    proximity_penalty_factor = st.number_input(
        "Kara za odległość między punktami (współczynnik)",
        min_value=0.0,
        value=0.2,
        step=0.05,
    )

with col_p6:
    long_jump_penalty = st.number_input(
        "Kara za przejazd >1.5h",
        min_value=0,
        value=10000,
        step=1000,
    )

time_limit_s = st.slider("Limit czasu szukania rozwiązania (sek.)", 2, 60, 12, 1)

service_time_s = int(service_time_h * 3600)
max_route_work_s = int(max_route_hours * 3600)
vehicle_fixed_cost = int(vehicle_fixed_cost_ui)
vehicle_fixed_cost = int(vehicle_fixed_cost_ui)
max_stops_per_route = int(max_stops_per_route_ui)
# =========================
# Helpers: wczytywanie
# =========================
def _norm_colname(c: str) -> str:
    return str(c).strip()

def normalize_postcode(code: str) -> str:
    s = str(code).strip().replace(" ", "")
    if len(s) == 5 and s.isdigit():
        return f"{s[:2]}-{s[2:]}"
    return s

def load_points(file) -> pd.DataFrame:
    name = (file.name or "").lower()
    if name.endswith(".xlsx"):
        df = pd.read_excel(file, dtype={"Kod": "string"})
    else:
        df = pd.read_csv(file, sep=";", encoding="utf-8", dtype={"Kod": "string"})
    df.columns = [_norm_colname(c) for c in df.columns]
    return df


def load_vehicles(file) -> pd.DataFrame:
    name = (file.name or "").lower()
    if name.endswith(".xlsx"):
        df = pd.read_excel(file)
    else:
        raw = file.getvalue().decode("utf-8", errors="ignore")
        if "\t" in raw and raw.count("\t") >= raw.count(";"):
            sep = "\t"
        elif ";" in raw:
            sep = ";"
        else:
            sep = ","
        from io import StringIO
        df = pd.read_csv(StringIO(raw), sep=sep)

    df.columns = [_norm_colname(c) for c in df.columns]
    return df


def safe_int(x, default=0) -> int:
    try:
        if pd.isna(x):
            return default
        s = str(x).strip().replace(",", ".")
        if s == "":
            return default
        return int(float(s))
    except Exception:
        return default


# =========================
# Helpers: Google APIs
# =========================
def geocode_cache_key(address: str) -> str:
    """Generuj klucz cache'u dla adresu"""
    raw = address.lower().strip().encode("utf-8")
    return hashlib.sha256(raw).hexdigest()


def geocode_address(address: str):
    cache_key = geocode_cache_key(address)
    cache_file = os.path.join(GEOCODING_CACHE_DIR, f"{cache_key}.json")

    # 1. Lokalny JSON cache
    if os.path.exists(cache_file):
        with open(cache_file, "r", encoding="utf-8") as f:
            cached = json.load(f)
        return cached["lat"], cached["lng"], cached["formatted"], cached["status"], "lokalny JSON"

    # 2. CSV cache z GitHub (session_state)
    if address in st.session_state["geocoding_cache_df"]["address"].values:
        row = st.session_state["geocoding_cache_df"][
            st.session_state["geocoding_cache_df"]["address"] == address
        ].iloc[0]
        cache_data = {
            "lat": row["lat"],
            "lng": row["lng"],
            "formatted": row["formatted_address"],
            "status": row["status"]
        }
        with open(cache_file, "w", encoding="utf-8") as f:
            json.dump(cache_data, f)
        return row["lat"], row["lng"], row["formatted_address"], row["status"], "GitHub CSV"

    # 3. Google Geocoding API
    url = "https://maps.googleapis.com/maps/api/geocode/json"
    params = {"address": address, "key": API_KEY}
    r = requests.get(url, params=params, timeout=25)
    data = r.json()

    if data.get("status") == "OK" and data.get("results"):
        loc = data["results"][0]["geometry"]["location"]
        formatted = data["results"][0].get("formatted_address", "")
        lat, lng, status = loc["lat"], loc["lng"], "OK"
    else:
        lat, lng, formatted, status = None, None, "", data.get("status", "UNKNOWN")

    cache_data = {"lat": lat, "lng": lng, "formatted": formatted, "status": status}
    with open(cache_file, "w", encoding="utf-8") as f:
        json.dump(cache_data, f)

    new_row = pd.DataFrame([{
        "address": address,
        "lat": lat,
        "lng": lng,
        "formatted_address": formatted,
        "status": status,
        "cached_at": datetime.now().isoformat()
    }])
    st.session_state["geocoding_cache_df"] = pd.concat(
        [st.session_state["geocoding_cache_df"], new_row],
        ignore_index=True
    )
    st.session_state["geocoding_updates"].add(address)

    return lat, lng, formatted, status, "Google API"


def save_geocoding_to_csv():
    """Zapisz cache geocodingu do CSV"""
    st.session_state["geocoding_cache_df"].to_csv(GEOCODING_CSV_PATH, index=False)
    return True


def update_geocoding_csv_github():
    """Commitnij zaktualizowany CSV do GitHub"""
    if not GITHUB_AVAILABLE:
        return False, "PyGithub nie zainstalowany"
    
    if "GITHUB_TOKEN" not in st.secrets:
        return False, "GITHUB_TOKEN nie znaleziony w secrets"
    
    if "GITHUB_REPO" not in st.secrets:
        return False, "GITHUB_REPO nie znaleziony w secrets (format: user/repo)"
    
    try:
        # Zapamiętaj liczbę zmian
        num_updates = len(st.session_state['geocoding_updates'])
        
        # Zaloguj do GitHub
        g = Github(st.secrets["GITHUB_TOKEN"])
        repo = g.get_repo(st.secrets["GITHUB_REPO"])
        
        # Przygotuj zawartość CSV
        csv_content = st.session_state["geocoding_cache_df"].to_csv(index=False)
        
        try:
            # Pobierz istniejący plik
            file = repo.get_contents(GEOCODING_CSV_PATH)
            # Update istniejącego pliku
            repo.update_file(
                GEOCODING_CSV_PATH,
                f"Auto: Update geocoding cache ({num_updates} nowych adresów)",
                csv_content,
                file.sha
            )
        except:
            # Utwórz nowy plik
            repo.create_file(
                GEOCODING_CSV_PATH,
                "Auto: Create geocoding cache",
                csv_content
            )
        
        # Zapisz lokalnie
        save_geocoding_to_csv()
        
        # Wyczyść tracker zmian
        st.session_state["geocoding_updates"].clear()
        
        return True, f"✅ Commitnięto {num_updates} nowych adresów do GitHub"
    
    
    except Exception as e:
        return False, f"❌ Błąd commitowania: {str(e)}"


def dm_pair_key(lat_o: float, lng_o: float, lat_d: float, lng_d: float) -> str:
    raw = f"{lat_o:.7f},{lng_o:.7f}|{lat_d:.7f},{lng_d:.7f}".encode()
    return hashlib.sha256(raw).hexdigest()


def push_pairs_cache_to_github(pairs: dict, n_new: int):
    if not GITHUB_AVAILABLE:
        return
    if "GITHUB_TOKEN" not in st.secrets or "GITHUB_REPO" not in st.secrets:
        return
    try:
        g = Github(st.secrets["GITHUB_TOKEN"])
        repo = g.get_repo(st.secrets["GITHUB_REPO"])
        content = json.dumps(pairs)
        try:
            existing = repo.get_contents(DM_PAIRS_CACHE_PATH)
            repo.update_file(
                DM_PAIRS_CACHE_PATH,
                f"Auto: Update DM pairs cache (+{n_new} par)",
                content, existing.sha,
            )
        except Exception:
            repo.create_file(DM_PAIRS_CACHE_PATH, "Auto: Create DM pairs cache", content)
        with open(DM_PAIRS_CACHE_PATH, "w", encoding="utf-8") as f:
            f.write(content)
    except Exception:
        pass


def format_latlng(lat, lng) -> str:
    return f"{lat},{lng}"


def chunked_idx(n, chunk_size):
    for i in range(0, n, chunk_size):
        yield list(range(i, min(i + chunk_size, n)))


def distance_matrix_google(origins, destinations, mode="driving"):
    url = "https://maps.googleapis.com/maps/api/distancematrix/json"
    params = {
        "origins": "|".join(origins),
        "destinations": "|".join(destinations),
        "mode": mode,
        "units": "metric",
        "key": API_KEY,
    }
    r = requests.get(url, params=params, timeout=35)
    data = r.json()

    if data.get("status") != "OK":
        raise RuntimeError(f"Distance Matrix API error: {data.get('status')} / {data.get('error_message')}")

    rows = data.get("rows", [])
    dist_m = [[None] * len(destinations) for _ in range(len(origins))]
    dur_s = [[None] * len(destinations) for _ in range(len(origins))]

    for i, row in enumerate(rows):
        elems = row.get("elements", [])
        for j, el in enumerate(elems):
            if el.get("status") == "OK":
                dist_m[i][j] = el["distance"]["value"]   # meters
                dur_s[i][j] = el["duration"]["value"]    # seconds
            else:
                dist_m[i][j] = None
                dur_s[i][j] = None

    return dist_m, dur_s


def build_full_matrix(points_latlng, mode="driving", sleep_s=0.05):
    n = len(points_latlng)
    coords = [tuple(map(float, p.split(","))) for p in points_latlng]
    pairs = st.session_state.get("dm_pairs_cache", {})

    missing_count = sum(
        1 for i in range(n) for j in range(n)
        if i != j and dm_pair_key(*coords[i], *coords[j]) not in pairs
    )

    if missing_count == 0:
        st.info(f"Macierz: wszystkie {n*(n-1)} par w cache — bez zapytań do Google API")
        dist = [[0]*n for _ in range(n)]
        dur  = [[0]*n for _ in range(n)]
        for i in range(n):
            for j in range(n):
                if i == j:
                    continue
                e = pairs[dm_pair_key(*coords[i], *coords[j])]
                dist[i][j] = e.get("d")
                dur[i][j]  = e.get("t")
        return dist, dur, True

    st.warning(f"Macierz: brakuje {missing_count} par w cache — pobieranie z Google Distance Matrix API")

    for batch_size in [10, 8, 5, 4, 2]:
        try:
            origin_batches = list(chunked_idx(n, batch_size))
            dest_batches   = list(chunked_idx(n, batch_size))

            batches_to_fetch = [
                (ob, db)
                for ob in origin_batches
                for db in dest_batches
                if any(
                    i != j and dm_pair_key(*coords[i], *coords[j]) not in pairs
                    for i in ob for j in db
                )
            ]

            new_pairs   = {}
            total_calls = len(batches_to_fetch)
            pb  = st.progress(0)
            txt = st.empty()

            for call_no, (ob, db) in enumerate(batches_to_fetch, 1):
                origins      = [points_latlng[i] for i in ob]
                destinations = [points_latlng[j] for j in db]
                dist_m, dur_s = distance_matrix_google(origins, destinations, mode=mode)

                for oi, i in enumerate(ob):
                    for dj, j in enumerate(db):
                        if i != j:
                            new_pairs[dm_pair_key(*coords[i], *coords[j])] = {
                                "d": dist_m[oi][dj],
                                "t": dur_s[oi][dj],
                            }

                pb.progress(int(call_no / total_calls * 100))
                txt.text(f"Macierz: {call_no}/{total_calls} zapytań (batch={batch_size})")
                time.sleep(sleep_s)

            pb.empty()
            txt.empty()

            pairs.update(new_pairs)
            st.session_state["dm_pairs_cache"] = pairs
            push_pairs_cache_to_github(pairs, len(new_pairs))
            st.success(f"Macierz: pobrano {len(new_pairs)} nowych par, cache zaktualizowany")

            dist = [[0]*n for _ in range(n)]
            dur  = [[0]*n for _ in range(n)]
            for i in range(n):
                for j in range(n):
                    if i == j:
                        continue
                    e = pairs.get(dm_pair_key(*coords[i], *coords[j])) or {}
                    dist[i][j] = e.get("d")
                    dur[i][j]  = e.get("t")
            return dist, dur, False

        except RuntimeError as e:
            if "MAX_ELEMENTS_EXCEEDED" in str(e):
                st.warning(f"MAX_ELEMENTS_EXCEEDED dla batch={batch_size}. Zmniejszam batch…")
                continue
            raise

    raise RuntimeError("Nie udało się pobrać macierzy – spróbuj mniejszej liczby punktów.")


# =========================
# OR-Tools: VRP
# =========================
def solve_vrp_capacity(
    cost_matrix,
    duration_matrix_s,
    demands,
    vehicle_capacities,
    depot=0,
    time_limit_s=60,
    drop_penalty=1_000_000,
    vehicle_fixed_cost=10_000,
    service_time_s=5400,
    max_route_work_s=28800,
    max_stops_per_route=5,
    proximity_penalty_factor=0.2
):


    n = len(cost_matrix)
    manager = pywrapcp.RoutingIndexManager(n, len(vehicle_capacities), depot)
    routing = pywrapcp.RoutingModel(manager)

    # koszt optymalizacji
    def cost_cb(from_index, to_index):
        frm = manager.IndexToNode(from_index)
        to = manager.IndexToNode(to_index)

        base_cost = int(cost_matrix[frm][to] or 0)
        extra_penalty = 0

        if frm != depot and to != depot:
            travel_s = int(duration_matrix_s[frm][to] or 0)

            # kara proporcjonalna (Twoja)
            extra_penalty += int(travel_s * proximity_penalty_factor)

            # 🔥 kara skokowa
            if travel_s > 5400:
                extra_penalty += long_jump_penalty

        return base_cost + extra_penalty

    cost_callback_index = routing.RegisterTransitCallback(cost_cb)
    routing.SetArcCostEvaluatorOfAllVehicles(cost_callback_index)

    # koszt użycia pojazdu
    for v in range(len(vehicle_capacities)):
        routing.SetFixedCostOfVehicle(int(vehicle_fixed_cost), v)

    # pojemność
    def demand_cb(from_index):
        node = manager.IndexToNode(from_index)
        return int(demands[node])

    demand_callback_index = routing.RegisterUnaryTransitCallback(demand_cb)
    routing.AddDimensionWithVehicleCapacity(
        demand_callback_index,
        0,
        [int(c) for c in vehicle_capacities],
        True,
        "Capacity",
    )

    # max 5 punktów na trasie (bez bazy)
    def visit_cb(from_index):
        node = manager.IndexToNode(from_index)
        return 0 if node == depot else 1

    visit_callback_index = routing.RegisterUnaryTransitCallback(visit_cb)
    routing.AddDimensionWithVehicleCapacity(
        visit_callback_index,
        0,
        [int(max_stops_per_route)] * len(vehicle_capacities),
        True,
        "Stops",
    )

    # czas pracy "w środku trasy"
    # liczymy:
    # - punkt -> punkt: przejazd + rozładunek w punkcie źródłowym
    # nie liczymy:
    # - baza -> pierwszy punkt
    # - ostatni punkt -> baza
    def work_time_cb(from_index, to_index):
        frm = manager.IndexToNode(from_index)
        to = manager.IndexToNode(to_index)

        # baza -> pierwszy punkt
        # nie liczymy do limitu 8h
        if frm == depot and to != depot:
            return 0

        # ostatni punkt -> baza
        # liczymy tylko rozładunek w ostatnim punkcie
        if frm != depot and to == depot:
            return int(service_time_s)

        # baza -> baza
        if frm == depot and to == depot:
            return 0

        # punkt -> punkt
        # liczymy przejazd + rozładunek w punkcie źródłowym
        travel_s = int(duration_matrix_s[frm][to] or 0)
        return int(travel_s + service_time_s)

    work_time_callback_index = routing.RegisterTransitCallback(work_time_cb)
    routing.AddDimension(
        work_time_callback_index,
        0,
        int(max_route_work_s),
        True,
        "WorkTime",
    )

    # pomijanie punktów za karę
    for node in range(1, n):
        routing.AddDisjunction([manager.NodeToIndex(node)], int(drop_penalty))

    search_params = pywrapcp.DefaultRoutingSearchParameters()
    search_params.first_solution_strategy = routing_enums_pb2.FirstSolutionStrategy.PARALLEL_CHEAPEST_INSERTION
    search_params.local_search_metaheuristic = routing_enums_pb2.LocalSearchMetaheuristic.GUIDED_LOCAL_SEARCH
    search_params.time_limit.FromSeconds(int(time_limit_s))
    search_params.log_search = True

    solution = routing.SolveWithParameters(search_params)
    if solution is None:
        return False, [], 0, []

    routes = []
    dropped = []

    for node in range(1, n):
        idx = manager.NodeToIndex(node)
        if solution.Value(routing.NextVar(idx)) == idx:
            dropped.append(node)

    for v in range(len(vehicle_capacities)):
        idx = routing.Start(v)
        route = [manager.IndexToNode(idx)]
        while not routing.IsEnd(idx):
            idx = solution.Value(routing.NextVar(idx))
            route.append(manager.IndexToNode(idx))
        routes.append(route)

    return True, routes, int(solution.ObjectiveValue()), dropped


# =========================
# Pipeline: przetwarzanie
# =========================
def ensure_ready_inputs():
    if points_file is None:
        st.error("Wgraj plik z punktami.")
        st.stop()
    if vehicles_file is None:
        st.error("Wgraj plik z pojazdami.")
        st.stop()


def build_nodes(points_df: pd.DataFrame) -> pd.DataFrame:
    ok_df = points_df[points_df["geocode_status"] == "OK"].copy()
    if ok_df.empty:
        raise RuntimeError("Brak punktów z poprawnym geokodowaniem (geocode_status != OK).")

    nodes = pd.DataFrame([{
        "node": "BASE",
        "name": BASE_NAME,
        "address": BASE_ADDRESS,
        "latitude": BASE_LAT,
        "longitude": BASE_LNG,
        "demand_wozki": 0,
    }])

    pts = ok_df.copy().reset_index(drop=True)
    pts["node"] = [f"P{i+1}" for i in range(len(pts))]
    pts["name"] = pts["Skrot_kontrahenta"].astype("string").fillna("").replace("nan", "").str.strip()

    fallback_name = (
        pts["Kod"].astype("string").fillna("").str.strip() + " " +
        pts["Miejscowosc"].astype("string").fillna("").str.strip()
    ).str.strip()
    pts["name"] = pts["name"].where(pts["name"].str.len() > 0, fallback_name)

    pts["address"] = pts["adres"].astype("string").fillna("").str.strip()
    pts["google_formatted_address"] = pts.get("formatted_address", "").astype("string").fillna("").str.strip()

    pts = pts[["node", "name", "address", "google_formatted_address", "latitude", "longitude", "demand_wozki"]]
    nodes = pd.concat([nodes, pts], ignore_index=True)
    return nodes


def fmt_hhmm(dt: datetime) -> str:
    return dt.strftime("%H:%M")


def calc_arrival_departure_for_route(route, duration_matrix_s, service_time_s):
    first_point_arrival_dt = datetime(2024, 1, 1, DAY_START_HOUR, 0, 0)

    times = []

    # trasa pusta albo tylko baza
    if len(route) == 0:
        return times

    # jeśli jest tylko baza -> baza 08:00
    if len(route) == 1:
        return [(first_point_arrival_dt, first_point_arrival_dt)]

    # czas dojazdu z bazy do 1. punktu
    first_travel_s = int(duration_matrix_s[route[0]][route[1]] or 0)

    # wyjazd z bazy wcześniej tak, aby w 1. punkcie być o 08:00
    base_departure_dt = first_point_arrival_dt - timedelta(seconds=first_travel_s)

    # baza
    times.append((base_departure_dt, base_departure_dt))

    current_dt = base_departure_dt

    for stop_no in range(1, len(route)):
        prev_node = route[stop_no - 1]
        node_idx = route[stop_no]

        travel_s = int(duration_matrix_s[prev_node][node_idx] or 0)
        arrival = current_dt + timedelta(seconds=travel_s)

        if node_idx == 0:
            # powrót do bazy - bez rozładunku
            departure = arrival
        else:
            # punkt klienta - rozładunek 1.5h
            departure = arrival + timedelta(seconds=service_time_s)

        times.append((arrival, departure))
        current_dt = departure

    return times


def calculate_route_metrics(routes, dist_km_df, dur_s_matrix):
    total_km = 0
    total_time_s = 0
    inter_point_times = []

    for route in routes:
        if len(route) < 2:
            continue

        for i in range(len(route) - 1):
            frm = route[i]
            to = route[i + 1]

            dist = dist_km_df.iloc[frm, to] or 0
            dur = dur_s_matrix[frm][to] or 0

            total_km += dist
            total_time_s += dur

            # tylko między punktami (bez bazy)
            if frm != 0 and to != 0:
                inter_point_times.append(dur)

    # mediana
    median_inter = 0
    if inter_point_times:
        inter_point_times_sorted = sorted(inter_point_times)
        n = len(inter_point_times_sorted)
        mid = n // 2
        if n % 2 == 0:
            median_inter = (inter_point_times_sorted[mid - 1] + inter_point_times_sorted[mid]) / 2
        else:
            median_inter = inter_point_times_sorted[mid]

    return {
        "total_km": total_km,
        "total_time_s": total_time_s,
        "total_inter_time_s": sum(inter_point_times),
        "median_inter_s": median_inter,
    }

def render_routes(
    routes: List[List[int]],
    nodes: pd.DataFrame,
    vehicle_ids: List[str],
    vehicle_caps: List[int],
    dur_s_matrix,
    service_time_s,
):
    node_names = nodes["name"].tolist()
    node_addr = nodes["address"].tolist()
    node_dem = nodes["demand_wozki"].astype(int).tolist()

    any_shown = False
    for v_idx, route in enumerate(routes):
        if len(route) <= 2:
            continue

        any_shown = True
        vehicle_label = vehicle_ids[v_idx] if v_idx < len(vehicle_ids) else str(v_idx + 1)
        cap = vehicle_caps[v_idx] if v_idx < len(vehicle_caps) else None

        used_capacity = sum(node_dem[idx] for idx in route if idx != 0)

        st.markdown(
            f"### Samochód {vehicle_label}, pojemność {cap} "
            f"(wykorzystane: {used_capacity})"
        )

        route_times = calc_arrival_departure_for_route(
            route,
            dur_s_matrix,
            service_time_s
        )

        rows = []
        for stop_no, node_idx in enumerate(route):
            arrival, departure = route_times[stop_no]

            rows.append({
                "Numer przystanku": stop_no,
                "Nazwa": node_names[node_idx],
                "Adres": node_addr[node_idx],
                "Ilość wózków": node_dem[node_idx],
                "Godzina przyjazdu i wyjazdu": f"{fmt_hhmm(arrival)} - {fmt_hhmm(departure)}",
            })

        st.dataframe(pd.DataFrame(rows).astype(str), use_container_width=True)

    if not any_shown:
        st.info("Solver nie przydzielił żadnych punktów do tras (wszystkie pojazdy BASE→BASE).")

def df_to_xlsx_bytes(df: pd.DataFrame, sheet_name: str = "Dane") -> bytes:
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)
    return output.getvalue()


_VEHICLE_COLORS = [
    "FFC000",  # amber
    "4472C4",  # blue
    "00B050",  # green
    "FF4B4B",  # red
    "7030A0",  # purple
    "00B0F0",  # sky blue
    "FF6600",  # orange
    "70AD47",  # light green
    "C00000",  # dark red
    "0070C0",  # dark blue
]


def routes_to_styled_xlsx_bytes(routes, nodes, vehicle_ids, vehicle_caps) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import PatternFill, Font, Border, Side, Alignment

    wb = Workbook()
    ws = wb.active
    ws.title = "Trasy"

    columns = ["nr pojazdu", "Pojemnosc", "Nazwa", "adres", "liczba wózków", "przewoźnik"]
    ws.append(columns)

    header_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    bold_font = Font(bold=True)
    for col in range(1, len(columns) + 1):
        cell = ws.cell(row=1, column=col)
        cell.fill = header_fill
        cell.font = bold_font

    node_names = nodes["name"].tolist()
    node_addr = nodes["address"].tolist()
    node_dem = nodes["demand_wozki"].astype(int).tolist()

    summary_fill = PatternFill(start_color="C0C0C0", end_color="C0C0C0", fill_type="solid")
    thick_bottom = Border(bottom=Side(style="thick", color="000000"))
    right_align = Alignment(horizontal="right")

    for v_idx, route in enumerate(routes):
        if len(route) <= 2:
            continue

        veh = vehicle_ids[v_idx] if v_idx < len(vehicle_ids) else str(v_idx + 1)
        cap = vehicle_caps[v_idx] if v_idx < len(vehicle_caps) else None
        total_wozki = 0
        first_stop = True

        color_hex = _VEHICLE_COLORS[v_idx % len(_VEHICLE_COLORS)]
        vehicle_fill = PatternFill(start_color=color_hex, end_color=color_hex, fill_type="solid")

        for node_idx in route:
            if node_idx == 0:
                continue
            total_wozki += node_dem[node_idx]
            ws.append([
                veh if first_stop else "",
                cap if first_stop else "",
                node_names[node_idx],
                node_addr[node_idx],
                node_dem[node_idx],
                "",
            ])
            ws.cell(row=ws.max_row, column=1).fill = vehicle_fill
            first_stop = False

        # szara linia podsumowująca — SUMA w kolumnie adres (4), wyrównana do prawej
        ws.append([veh, "", "", "SUMA", total_wozki, ""])
        summary_row_idx = ws.max_row
        for col in range(1, 6):
            cell = ws.cell(row=summary_row_idx, column=col)
            cell.fill = summary_fill
            cell.font = bold_font
            cell.border = thick_bottom
        ws.cell(row=summary_row_idx, column=4).alignment = right_align

    output = BytesIO()
    wb.save(output)
    return output.getvalue()


def routes_to_word_bytes(routes, nodes, vehicle_ids, vehicle_caps, dur_s_matrix, service_time_s) -> bytes:
    doc = Document()
    doc.add_heading("Planowanie tras Plantpol", level=1)

    node_names = nodes["name"].tolist()
    node_addr = nodes["address"].tolist()
    node_dem = nodes["demand_wozki"].astype(int).tolist()

    any_shown = False

    for v_idx, route in enumerate(routes):
        if len(route) <= 2:
            continue

        any_shown = True
        vehicle_label = vehicle_ids[v_idx] if v_idx < len(vehicle_ids) else str(v_idx + 1)
        cap = vehicle_caps[v_idx] if v_idx < len(vehicle_caps) else None
        used_capacity = sum(node_dem[idx] for idx in route if idx != 0)

        doc.add_heading(
            f"Samochód {vehicle_label}, pojemność {cap} (wykorzystane: {used_capacity})",
            level=2
        )

        route_times = calc_arrival_departure_for_route(route, dur_s_matrix, service_time_s)

        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        hdr = table.rows[0].cells
        hdr[0].text = "Numer przystanku"
        hdr[1].text = "Nazwa"
        hdr[2].text = "Adres"
        hdr[3].text = "Ilość wózków"
        hdr[4].text = "Godzina przyjazdu i wyjazdu"

        for stop_no, node_idx in enumerate(route):
            arrival, departure = route_times[stop_no]
            row = table.add_row().cells
            row[0].text = str(stop_no)
            row[1].text = str(node_names[node_idx])
            row[2].text = str(node_addr[node_idx])
            row[3].text = str(node_dem[node_idx])
            row[4].text = f"{fmt_hhmm(arrival)} - {fmt_hhmm(departure)}"

        doc.add_paragraph("")

    if not any_shown:
        doc.add_paragraph("Brak tras do wydruku.")

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


# =========================
# Mapa
# =========================
def build_map_html(routes, nodes, vehicle_ids) -> str:
    import folium

    lats = nodes["latitude"].tolist()
    lngs = nodes["longitude"].tolist()
    names = nodes["name"].tolist()
    addrs = nodes["address"].tolist()

    valid = [(la, ln) for la, ln in zip(lats, lngs) if la is not None and ln is not None]
    center_lat = sum(x[0] for x in valid) / len(valid)
    center_lng = sum(x[1] for x in valid) / len(valid)

    m = folium.Map(location=[center_lat, center_lng], zoom_start=7, tiles="OpenStreetMap")

    folium.Marker(
        location=[lats[0], lngs[0]],
        tooltip="Plantpol – baza",
        icon=folium.Icon(color="black", icon="home", prefix="fa"),
    ).add_to(m)

    for v_idx, route in enumerate(routes):
        if len(route) <= 2:
            continue

        veh = vehicle_ids[v_idx] if v_idx < len(vehicle_ids) else str(v_idx + 1)
        color = f"#{_VEHICLE_COLORS[v_idx % len(_VEHICLE_COLORS)]}"

        coords = [
            (lats[idx], lngs[idx]) for idx in route
            if lats[idx] is not None and lngs[idx] is not None
        ]
        if len(coords) >= 2:
            folium.PolyLine(
                coords,
                color=color,
                weight=2.5,
                opacity=0.8,
                tooltip=f"Pojazd {veh}",
            ).add_to(m)

        for node_idx in route:
            if node_idx == 0:
                continue
            if lats[node_idx] is None or lngs[node_idx] is None:
                continue
            folium.CircleMarker(
                location=[lats[node_idx], lngs[node_idx]],
                radius=7,
                color=color,
                fill=True,
                fill_color=color,
                fill_opacity=0.9,
                tooltip=f"Pojazd {veh} | {names[node_idx]} | {addrs[node_idx]}",
            ).add_to(m)

    return m.get_root().render()


# =========================
# Zakładki
# =========================
tab_result, tab_geocode, tab_matrix, tab_map = st.tabs(["Wynik trasowania", "Geokodowanie", "Macierz", "Mapa"])

with tab_result:
    run_all = st.button("URUCHOM CAŁY PROCES", type="primary", use_container_width=True)

    if run_all:



        ensure_ready_inputs()

        stage = st.empty()
        stage.info("Etap 0/3: wczytuję dane…")

        points_df = load_points(points_file)
        vehicles_df = load_vehicles(vehicles_file)

        required_points = ["Kod", "Miejscowosc", "Skrot_kontrahenta", "Ilość wózków"]
        missing = [c for c in required_points if c not in points_df.columns]
        if missing:
            st.error(f"Brakuje kolumn w pliku punktów: {', '.join(missing)}")
            st.stop()

        n_points = len(points_df)
        if n_points > MAX_POINTS:
            st.error(f"Za dużo punktów do testów: {n_points}. Limit to {MAX_POINTS}.")
            st.stop()

        required_veh = ["samochód", "ilość wózków"]
        missing_v = [c for c in required_veh if c not in vehicles_df.columns]
        if missing_v:
            st.error(f"Brakuje kolumn w pliku pojazdów: {', '.join(missing_v)}")
            st.stop()

        vehicles_df["samochód"] = vehicles_df["samochód"].astype("string").fillna("").str.strip()
        vehicles_df["ilość wózków"] = vehicles_df["ilość wózków"].apply(safe_int)
        vehicles_df = vehicles_df[vehicles_df["ilość wózków"] > 0].copy()

        if vehicles_df.empty:
            st.error("Plik pojazdów nie zawiera żadnych pojazdów z dodatnią pojemnością.")
            st.stop()

        vehicle_ids = vehicles_df["samochód"].tolist()
        vehicle_caps = vehicles_df["ilość wózków"].astype(int).tolist()

        points_df["Kod"] = points_df["Kod"].astype("string").fillna("").str.strip().apply(normalize_postcode)
        points_df["Miejscowosc"] = points_df["Miejscowosc"].astype("string").fillna("").str.strip()

        points_df["adres"] = (
            points_df["Kod"].where(points_df["Kod"] != "", "") + " " +
            points_df["Miejscowosc"].where(points_df["Miejscowosc"] != "", "")
        ).str.strip() + ", Polska"
        points_df["demand_wozki"] = points_df["Ilość wózków"].apply(safe_int)

        stage.info("Etap 1/3: geokodowanie…")
        geo_pb = st.progress(0)
        geo_txt = st.empty()

        lats, lngs, formatted, statuses = [], [], [], []
        geo_sources = {"lokalny JSON": 0, "GitHub CSV": 0, "Google API": 0}
        total = len(points_df)
        for i, addr in enumerate(points_df["adres"]):
            lat, lng, fmt, status, source = geocode_address(addr)
            lats.append(lat)
            lngs.append(lng)
            formatted.append(fmt)
            statuses.append(status)
            geo_sources[source] = geo_sources.get(source, 0) + 1

            p = int((i + 1) / total * 100)
            geo_pb.progress(p)
            geo_txt.text(f"Geokodowanie: {i+1}/{total} ({p}%) — ostatni: {source}")

        geo_pb.empty()
        geo_txt.empty()
        st.info(
            f"Geokodowanie — źródła: "
            f"{geo_sources['lokalny JSON']} z lokalnego cache | "
            f"{geo_sources['GitHub CSV']} z GitHub CSV | "
            f"{geo_sources['Google API']} nowych z Google API"
        )

        points_df["latitude"] = lats
        points_df["longitude"] = lngs
        points_df["formatted_address"] = formatted
        points_df["geocode_status"] = statuses

        st.session_state["points_df"] = points_df.copy()

        ok_cnt = int((points_df["geocode_status"] == "OK").sum())
        if ok_cnt == 0:
            st.error("Nie udało się geokodować żadnego punktu (OK=0).")
            st.stop()

        # Auto-commit nowych adresów do GitHub
        if len(st.session_state["geocoding_updates"]) > 0:
            success, msg = update_geocoding_csv_github()
            if success:
                st.toast(f"✅ {msg}")
            else:
                st.warning(f"⚠️ Nie udało się zsynchronizować z GitHub: {msg}")

        stage.info("Etap 2/3: pobieram macierz dystansów/czasów…")

        nodes = build_nodes(points_df)
        st.session_state["nodes_df"] = nodes.copy()

        points_latlng = [format_latlng(r.latitude, r.longitude) for r in nodes.itertuples(index=False)]

        dist_m, dur_s, from_cache = build_full_matrix(points_latlng, mode="driving", sleep_s=0.05)

        labels = nodes["node"].tolist()
        dist_km_df = pd.DataFrame(dist_m, index=labels, columns=labels) / 1000.0
        dur_min_df = pd.DataFrame(dur_s, index=labels, columns=labels) / 60.0

        st.session_state["dist_km_df"] = dist_km_df
        st.session_state["dur_min_df"] = dur_min_df
        st.session_state["matrix_from_cache"] = from_cache
        st.session_state["dur_s_matrix"] = dur_s

        stage.info("Etap 3/3: liczę trasy (OR-Tools)…")
        solver_status = st.empty()
        solver_pb = st.progress(0)

        # zawsze optymalizujemy po czasie
        cost = (dur_min_df.fillna(1e9) * 60.0).round().astype("int64").values.tolist()

        # osobna macierz czasu do constraintów czasowych
        duration_cost_s = (dur_min_df.fillna(1e9) * 60.0).round().astype("int64").values.tolist()

        demands = nodes["demand_wozki"].fillna(0).astype(int).tolist()

        if sum(demands) > sum(vehicle_caps):
            st.error(f"Suma wózków ({sum(demands)}) > łączna pojemność floty ({sum(vehicle_caps)}).")
            st.stop()

        solver_status.text("Solver: start…")
        solver_pb.progress(10)

        ok, routes, obj, dropped = solve_vrp_capacity(
            cost_matrix=cost,
            duration_matrix_s=duration_cost_s,
            demands=demands,
            vehicle_capacities=vehicle_caps,
            depot=0,
            time_limit_s=int(time_limit_s),
            vehicle_fixed_cost=vehicle_fixed_cost,
            service_time_s=service_time_s,
            max_route_work_s=max_route_work_s,
            max_stops_per_route=max_stops_per_route,
            proximity_penalty_factor=proximity_penalty_factor,
)

        solver_pb.progress(100)
        solver_status.empty()
        solver_pb.empty()
        stage.empty()

        if dropped:
            st.warning(f"Pominięte punkty: {len(dropped)} (sprawdź macierz/geo/limity dla tych punktów).")
            dropped_df = nodes.iloc[dropped][["node", "name", "address", "demand_wozki"]].copy()
            st.dataframe(dropped_df.astype(str), use_container_width=True)

        if not ok:
            st.error("Nie znaleziono rozwiązania w zadanym limicie czasu.")
            st.stop()

        st.session_state["routes"] = routes
        st.session_state["obj"] = obj
        st.session_state["vehicle_ids"] = vehicle_ids
        st.session_state["vehicle_caps"] = vehicle_caps
        service_time_s = int(service_time_h * 3600)
        max_route_work_s = int(max_route_hours * 3600)

    if "routes" in st.session_state and "nodes_df" in st.session_state:
        routes = st.session_state["routes"]
        nodes = st.session_state["nodes_df"]
        vehicle_ids = st.session_state.get("vehicle_ids", [])
        vehicle_caps = st.session_state.get("vehicle_caps", [])
        dur_s_matrix = st.session_state.get("dur_s_matrix")
        dist_km_df = st.session_state.get("dist_km_df")
      

        metrics = calculate_route_metrics(routes, dist_km_df, dur_s_matrix)

        def fmt_h(x):
            return f"{x/3600:.1f} h"

        col_m1, col_m2, col_m3, col_m4 = st.columns(4)

        col_m1.metric("Łączny dystans", f"{metrics['total_km']:.0f} km")
        col_m2.metric("Łączny czas jazdy", fmt_h(metrics["total_time_s"]))
        col_m3.metric("Czas między punktami", fmt_h(metrics["total_inter_time_s"]))
        col_m4.metric("Mediana między punktami", fmt_h(metrics["median_inter_s"]))


        render_routes(
            routes,
            nodes,
            vehicle_ids,
            vehicle_caps,
            dur_s_matrix,
            service_time_s
        )

        xlsx_bytes = routes_to_styled_xlsx_bytes(routes, nodes, vehicle_ids, vehicle_caps)
        st.download_button(
            "Pobierz wynik tras XLSX",
            data=xlsx_bytes,
            file_name="wynik_trasy.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

        # Word do wydruku
        word_bytes = routes_to_word_bytes(
            routes,
            nodes,
            vehicle_ids,
            vehicle_caps,
            dur_s_matrix,
            service_time_s
        )
        st.download_button(
            "Pobierz wynik tras Word",
            data=word_bytes,
            file_name="wynik_trasy.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


with tab_geocode:
    st.subheader("Geokodowanie – podgląd")
    if "points_df" in st.session_state:
        st.dataframe(st.session_state["points_df"].astype(str), use_container_width=True)
    else:
        st.info("Brak danych – uruchom proces w zakładce Wynik trasowania.")


with tab_matrix:
    st.subheader("Macierz – dystans i czas")
    if "dist_km_df" in st.session_state and "dur_min_df" in st.session_state and "nodes_df" in st.session_state:
        from_cache = st.session_state.get("matrix_from_cache", False)
        st.caption("Źródło: " + ("cache" if from_cache else "Google Distance Matrix API"))

        st.write("Węzły (baza + punkty OK):")
        st.dataframe(st.session_state["nodes_df"].astype(str), use_container_width=True)

        st.write("Dystans [km]:")
        st.dataframe(st.session_state["dist_km_df"].round(3), use_container_width=True)

        st.write("Czas [min]:")
        st.dataframe(st.session_state["dur_min_df"].round(1), use_container_width=True)
    else:
        st.info("Brak danych – uruchom proces w zakładce Wynik trasowania.")


with tab_map:
    st.subheader("Mapa tras")
    if "routes" in st.session_state and "nodes_df" in st.session_state:
        map_html = build_map_html(
            st.session_state["routes"],
            st.session_state["nodes_df"],
            st.session_state.get("vehicle_ids", []),
        )
        st.components.v1.html(map_html, height=600, scrolling=False)
        st.download_button(
            "Pobierz mapę HTML",
            data=map_html.encode("utf-8"),
            file_name="mapa_tras.html",
            mime="text/html",
            use_container_width=True,
        )
    else:
        st.info("Brak danych – uruchom proces w zakładce Wynik trasowania.")